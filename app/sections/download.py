"""Section for saving a report into .docx file."""

import os
import tempfile
import uuid
from typing import List, Union

import pandas as pd
import plotly.graph_objects as go
import plotly.io as pio
from backend.data_extractors import ExitEmbaExtractor
from backend.plotter import PlotGenerator
from docx import Document
from docx.shared import Inches
from plotly.colors import qualitative


def save_plot_as_image(fig: go.Figure) -> str:
    """Save a Plotly figure as a temporary PNG image file.

    Args:
        fig (go.Figure): The Plotly figure to save.

    Returns:
        str: Path to the saved temporary image file.
    """
    tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmpfile.close()
    pio.write_image(fig, tmpfile.name, format="png", width=800, height=600)
    return tmpfile.name


def create_all_plots(emba_extractor: ExitEmbaExtractor) -> List[go.Figure]:
    """Generate all survey-related visualizations as Plotly figures.

    Args:
        emba_extractor (ExitEmbaExtractor): Extractor providing necessary data.

    Returns:
        List[go.Figure]: List of Plotly figures representing the generated plots.
    """

    def _create_pie(title: str, data: Union[pd.Dataframe, pd.Series]) -> go.Figure:
        """Create a pie chart with qualitative color palette.

        Args:
            title (str): Title of the plot.
            data (Union[pd.Dataframe, pd.Series]): Data to plot.

        Returns:
            go.Figure: Plotly pie chart figure.
        """
        fig = PlotGenerator.make_piechart(data=data, title=title)

        labels = fig.data[0].labels
        color_list = qualitative.Set1 + qualitative.Set2 + qualitative.Set3
        fig.data[0].marker.colors = color_list[: len(labels)]
        return fig

    def _create_boxplot(title: str, data: Union[pd.Dataframe, pd.Series]) -> go.Figure:
        """Create a colored boxplot.

        Args:
            title (str): Title of the plot.
            data (Union[pd.Dataframe, pd.Series]): Data to plot.

        Returns:
            go.Figure: Plotly boxplot figure.
        """
        fig = PlotGenerator.make_boxplot(data=data, title=title)
        for i, trace in enumerate(fig.data):
            trace.marker.color = qualitative.D3[i % len(qualitative.D3)]
            trace.line.color = trace.marker.color
        return fig

    def _create_nps(
        title: str, barnames: list[str], values: list[float], colors: list[str]
    ) -> go.Figure:
        """Create a Net Promoter Score (NPS) bar plot.

        Args:
            title (str): Plot title.
            barnames (list[str]): Labels for bars.
            values (list[float]): Values for each bar.
            colors (list[str]): Bar colors.

        Returns:
            go.Figure: Plotly barplot figure.
        """
        fig = PlotGenerator.make_nps_barplot(
            barnames=barnames, values=values, colors=colors, title=title
        )
        return fig

    def _create_stacked_bar(
        title: str, data: Union[pd.Dataframe, pd.Series]
    ) -> go.Figure:
        """Create a stacked bar plot for scores.

        Args:
            title (str): Plot title.
            data (Union[pd.Dataframe, pd.Series]): Score distribution data.

        Returns:
            go.Figure: Plotly stacked bar chart.
        """
        fig = PlotGenerator.make_stacked_barplot(
            score_counts=data,
            title=title,
            x_axs_title="Предмет",
            y_axs_title="Количество студентов",
        )
        return fig

    def _create_hbar(
        title: str,
        data: Union[pd.Dataframe, pd.Series],
        x_axs_title: str,
        y_axs_title: str,
    ) -> go.Figure:
        """Create a horizontal bar plot.

        Args:
            title (str): Plot title.
            data (Union[pd.Dataframe, pd.Series]): Data for plotting.
            x_axs_title (str): Label for X-axis.
            y_axs_title (str): Label for Y-axis.

        Returns:
            go.Figure: Plotly horizontal bar chart.
        """
        fig = PlotGenerator.make_hbarplot(
            score_counts=data,
            title=title,
            x_axs_title=x_axs_title,
            y_axs_title=y_axs_title,
        )
        return fig

    figs = []
    figs.append(_create_pie("Распределение возрастов", emba_extractor.get_age_distr()))
    figs.append(
        _create_pie(
            "Распределение студентов по индустриям", emba_extractor.get_industry_distr()
        )
    )
    figs.append(
        _create_pie(
            "Распределение студентов по занимаемым должностям",
            emba_extractor.get_job_distr(),
        )
    )
    figs.append(
        _create_boxplot("Оценка программы в целом", emba_extractor.get_basic_csi())
    )
    figs.append(
        _create_boxplot(
            "Распределение оценок различных составляющих курса",
            emba_extractor.get_overall_csi(),
        )
    )
    figs.append(
        _create_boxplot("Оценка дизайна программы", emba_extractor.get_design_csi())
    )
    figs.append(
        _create_boxplot(
            "Оценка международных модулей", emba_extractor.get_intern_modules_csi()
        )
    )
    figs.append(
        _create_boxplot("Оценка работы команды курса", emba_extractor.get_support_csi())
    )
    figs.append(
        _create_boxplot("Оценка качества группы", emba_extractor.get_group_csi())
    )
    barnames, values, colors = emba_extractor.get_industry_nps()
    figs.append(_create_nps("NPS EMBA-35 vs NPS индустрии", barnames, values, colors))
    barnames, values, colors = emba_extractor.get_programs_nps()
    figs.append(
        _create_nps("Сравнение NPS разных программ Школы", barnames, values, colors)
    )
    figs.append(
        _create_stacked_bar("Распределение оценок PILOs", emba_extractor.get_pilos())
    )
    figs.append(
        _create_hbar(
            "Лучшие преподаватели",
            emba_extractor.get_top_lectors(),
            x_axs_title="Количество голосов",
            y_axs_title="Преподаватель",
        )
    )
    figs.append(
        _create_hbar(
            "В каких активностях готовы участвовать выпускники",
            emba_extractor.get_collaborators(),
            x_axs_title="Количество согласившихся",
            y_axs_title="Готов участвовать в...",
        )
    )

    return figs


def generate_docx(emba_extractor: ExitEmbaExtractor) -> str:
    """Generate a DOCX report from all generated plots.

    Args:
        emba_extractor (ExitEmbaExtractor): Data extractor used to populate figures.

    Returns:
        str: File path to the saved DOCX report.
    """
    doc = Document()
    doc.add_heading("Сгенерированный отчет по EMBA-35", 0)
    for plot in create_all_plots(emba_extractor):
        img_path = save_plot_as_image(plot)
        doc.add_picture(img_path, width=Inches(5.5))
        os.remove(img_path)

    # Save document
    file_path = f"report_{uuid.uuid4().hex[:6]}.docx"
    doc.save(file_path)
    return file_path
